"""Word (.docx) loader → common Document model."""

from __future__ import annotations

import logging
from pathlib import Path

from backend.loaders.base import ChapterHeading, Document

logger = logging.getLogger(__name__)


def load_word(filepath: str | Path) -> Document:
    """Parse a .docx file into the common Document model.

    Chapter detection based on built-in heading styles.
    """
    filepath = Path(filepath)

    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx is required for Word loading: pip install python-docx")

    docx = DocxDocument(str(filepath))
    paragraphs: list[str] = []
    chapters: list[ChapterHeading] = []

    for para in docx.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        paragraphs.append(text)

        # Detect headings by style name
        style_name = para.style.name.lower() if para.style else ""
        if style_name.startswith("heading") or style_name.startswith("heading "):
            level_str = style_name.replace("heading", "").strip()
            try:
                level = int(level_str) if level_str else 1
            except ValueError:
                level = 1
            chapters.append(ChapterHeading(
                title=text,
                level=min(level, 3),
                page=len(paragraphs),
                order=len(chapters),
            ))

    full_text = "\n".join(paragraphs)

    return Document(
        filename=filepath.name,
        file_type="docx",
        chapters=chapters,
        full_text=full_text,
        char_count=len(full_text),
    )
